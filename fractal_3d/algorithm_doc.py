"""Generate a bilingual (RU/UZ) Word document describing the DCNF-CRF depth
algorithm: pipeline block-diagram, numbered steps with formulas, and the
evaluation metrics (Liu et al. CVPR 2015, Tables 1-3) with explanations.

Used by the /algorithm_doc API endpoint. Only python-docx + matplotlib (Agg).
"""
from __future__ import annotations

import io


# Numbered algorithm steps: (title_ru, title_uz, formula, explain_ru, explain_uz)
_STEPS = [
    ("SLIC сегментация", "SLIC segmentatsiya", "image → n суперпикселей",
     "Изображение разбивается на ~500 однородных регионов (суперпикселей).",
     "Tasvir ~500 bir jinsli mintaqaga (superpikselga) bo‘linadi."),
    ("Unary z_p", "Unary z_p", "U(y_p) = (y_p − z_p)²",
     "Грубая глубина каждого суперпикселя (DA V2 / pseudo-cues / shape-from-shading).",
     "Har bir superpikselning dastlabki chuqurligi (DA V2 / pseudo / shape-from-shading)."),
    ("Pairwise R_pq", "Pairwise R_pq",
     "S^(k)_pq = exp(−γ_k·‖s^(k)_p − s^(k)_q‖);  R_pq = Σ β_k·S^(k)_pq",
     "3 similarity: LAB color, color histogram, LBP texture — связывают похожие соседние регионы.",
     "3 similarity: LAB rang, rang gistogrammasi, LBP tekstura — o‘xshash qo‘shni mintaqalarni bog‘laydi."),
    ("Энергия CRF", "CRF energiyasi",
     "E(y) = Σ(y_p−z_p)² + Σ½·R_pq·(y_p−y_q)² = yᵀAy − 2zᵀy + zᵀz",
     "Сумма unary (близость к z) и pairwise (гладкость). A = I + D − R, D_pp = Σ_q R_pq.",
     "Unary (z ga yaqinlik) va pairwise (silliqlik) yig‘indisi. A = I + D − R."),
    ("MAP-решение", "MAP yechimi", "y* = A⁻¹ z",
     "Точное решение в замкнутой форме (минимум квадратичной энергии).",
     "Yopiq shakldagi aniq yechim (kvadratik energiya minimumi)."),
    ("Guided filter + blend", "Guided filter + blend", "де-блок + смешивание с unary",
     "Убираем ступеньки суперпикселей, сохраняя границы изображения; смешиваем с unary.",
     "Superpiksel zinapoyalarini olib tashlaymiz, tasvir chegaralarini saqlaymiz."),
    ("3D-реконструкция", "3D rekonstruksiya", "depth → marching cubes / heightmap → 3D",
     "Финальная карта глубины превращается в 3D-поверхность.",
     "Yakuniy chuqurlik xaritasi 3D yuzaga aylantiriladi."),
]

# Metrics: (symbol, name_ru, name_uz, formula, explain_ru, explain_uz, better)
_METRICS = [
    ("rel", "Средняя относительная ошибка", "O‘rtacha nisbiy xato",
     "rel = (1/T)·Σ |d_gt − d| / d_gt",
     "Отличие глубины от истинной, в долях; деление на истину делает близкие/дальние сравнимыми.",
     "Chuqurlikning haqiqiydan farqi, ulushlarda.", "меньше / kichikroq"),
    ("log10", "Средняя логарифмическая ошибка", "O‘rtacha logarifmik xato",
     "log10 = (1/T)·Σ |log₁₀(d_gt) − log₁₀(d)|",
     "Ошибка в лог-масштабе; подходит когда глубины меняются на порядки.",
     "Logarifmik masshtabdagi xato.", "меньше / kichikroq"),
    ("rms", "Среднеквадратичная ошибка", "O‘rtacha kvadratik xato (RMS)",
     "rms = √( (1/T)·Σ (d_gt − d)² )",
     "Корень из среднего квадрата ошибок; сильнее штрафует большие промахи.",
     "Xatolar kvadrati o‘rtachasidan ildiz.", "меньше / kichikroq"),
    ("δ < 1.25", "Точность с порогом", "Chegarali aniqlik",
     "% пикселей где max(d_gt/d, d/d_gt) < 1.25",
     "Доля пикселей в пределах 25% от истины; пороги 1.25, 1.25², 1.25³.",
     "25% ichidagi piksellar ulushi.", "больше / kattaroq"),
]


def _mono(run):
    run.font.name = "Consolas"
    from docx.shared import Pt
    run.font.size = Pt(10)


def _pipeline_image() -> io.BytesIO | None:
    """Render a simple DCNF-CRF pipeline block-diagram as a PNG (BytesIO)."""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

        fig, ax = plt.subplots(figsize=(7.2, 8.0))
        ax.set_xlim(0, 10); ax.set_ylim(0, 12); ax.axis("off")

        def box(x, y, w, h, text, fc="#e8f0fe"):
            ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.1",
                                        fc=fc, ec="#1a3a6a", lw=1.2))
            ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=9)

        def arrow(x0, y0, x1, y1):
            ax.add_patch(FancyArrowPatch((x0, y0), (x1, y1), arrowstyle="-|>",
                                         mutation_scale=14, color="#1a3a6a", lw=1.2))

        box(3, 11, 4, 0.8, "Входное изображение (RGB)")
        arrow(5, 11, 5, 10.5)
        box(2.5, 9.6, 5, 0.8, "SLIC суперпиксели (~500)")
        arrow(4, 9.6, 2.6, 8.7); arrow(6, 9.6, 7.4, 8.7)
        box(0.2, 7.4, 4.4, 1.3,
            "Unary:\nz_p = DA V2 / SfS / pseudo\nU = (y_p − z_p)²", fc="#e6f7ec")
        box(5.4, 7.4, 4.4, 1.3,
            "Pairwise:\ncolor · histogram · LBP\nR_pq = Σ β_k·S^(k)", fc="#fdeede")
        arrow(2.4, 7.4, 4.6, 6.6); arrow(7.6, 7.4, 5.4, 6.6)
        box(2.5, 5.7, 5, 0.9, "E(y) = Σ(y−z)² + Σ½R(y_p−y_q)²")
        arrow(5, 5.7, 5, 5.2)
        box(3.2, 4.3, 3.6, 0.8, "A = I + D − R")
        arrow(5, 4.3, 5, 3.8)
        box(3.2, 2.9, 3.6, 0.8, "MAP:  y* = A⁻¹ z", fc="#e6f7ec")
        arrow(5, 2.9, 5, 2.4)
        box(2.5, 1.5, 5, 0.8, "Guided filter + blend")
        arrow(5, 1.5, 5, 1.0)
        box(2.5, 0.1, 5, 0.8, "Карта глубины → 3D")

        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=300, bbox_inches="tight", facecolor="white")
        plt.close(fig)
        buf.seek(0)
        return buf
    except Exception:
        return None


def _gost_flowchart(nodes, arrows, figsize):
    """Render a GOST-style flowchart PNG (BytesIO).

    nodes: list of (shape, cx, cy, w, h, text) where shape in
      {start, io, proc, cond}. arrows: list of (x0, y0, x1, y1, label).
    """
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        from matplotlib.patches import FancyBboxPatch, FancyArrowPatch, Polygon, Ellipse

        fig, ax = plt.subplots(figsize=figsize)
        ax.set_xlim(0, 10)
        ymax = max(cy + h for _, _, cy, _, h, _ in nodes) + 1
        ax.set_ylim(0, ymax)
        ax.axis("off")

        for shape, cx, cy, w, h, text in nodes:
            if shape == "start":
                ax.add_patch(Ellipse((cx, cy), w, h, fc="#dbeafe", ec="#1e3a8a", lw=1.3))
            elif shape == "io":
                sk = w * 0.18
                ax.add_patch(Polygon([(cx - w / 2 + sk, cy - h / 2), (cx + w / 2, cy - h / 2),
                                      (cx + w / 2 - sk, cy + h / 2), (cx - w / 2, cy + h / 2)],
                                     closed=True, fc="#fef9c3", ec="#a16207", lw=1.2))
            elif shape == "cond":
                ax.add_patch(Polygon([(cx, cy + h / 2), (cx + w / 2, cy), (cx, cy - h / 2),
                                      (cx - w / 2, cy)], closed=True,
                                     fc="#fde9d8", ec="#b45309", lw=1.2))
            else:  # proc
                ax.add_patch(FancyBboxPatch((cx - w / 2, cy - h / 2), w, h,
                                            boxstyle="round,pad=0.02", fc="#e8f0fe",
                                            ec="#1e3a8a", lw=1.1))
            ax.text(cx, cy, text, ha="center", va="center", fontsize=8.2)

        for x0, y0, x1, y1, label in arrows:
            ax.add_patch(FancyArrowPatch((x0, y0), (x1, y1), arrowstyle="-|>",
                                         mutation_scale=12, color="#334155", lw=1.1,
                                         connectionstyle="arc3,rad=0"))
            if label:
                ax.text((x0 + x1) / 2 + 0.2, (y0 + y1) / 2, label, fontsize=8,
                        color="#b45309", ha="left", va="center")

        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=300, bbox_inches="tight", facecolor="white")
        plt.close(fig)
        buf.seek(0)
        return buf
    except Exception:
        return None


def _classic_depth_image():
    """GOST flowchart of the DCNF-CRF depth algorithm."""
    N = [
        ("start", 5, 15.2, 2.6, 0.8, "Начало / Boshlash"),
        ("io", 5, 14.0, 4.2, 0.8, "Ввод: изображение x"),
        ("proc", 5, 12.8, 4.6, 0.8, "SLIC → n суперпикселей"),
        ("proc", 5, 11.7, 2.0, 0.7, "p = 1"),
        ("cond", 5, 10.4, 2.6, 1.2, "p ≤ n ?"),
        ("proc", 5, 9.0, 3.2, 0.8, "z_p = DA_V2(p)"),
        ("proc", 5, 7.9, 2.4, 0.7, "p = p + 1"),
        ("proc", 5, 6.6, 4.8, 0.8, "R_pq = Σ β_k·S^(k)"),
        ("proc", 5, 5.5, 3.0, 0.7, "A = I + D − R"),
        ("cond", 5, 4.1, 2.8, 1.2, "A обратима?"),
        ("proc", 8.4, 4.1, 2.6, 0.8, "A ← A + εI"),
        ("proc", 5, 2.8, 3.6, 0.8, "y* = A⁻¹·z"),
        ("proc", 5, 1.8, 3.4, 0.7, "Guided filter + blend"),
        ("io", 5, 0.8, 4.2, 0.7, "Вывод: карта y → 3D"),
        ("start", 5, -0.3, 2.6, 0.8, "Конец / Tugadi"),
    ]
    A = [
        (5, 14.8, 5, 14.4, ""), (5, 13.6, 5, 13.2, ""), (5, 12.4, 5, 12.05, ""),
        (5, 11.35, 5, 11.0, ""), (5, 9.8, 5, 9.4, "да/ha"),
        (5, 8.6, 5, 8.25, ""), (3.8, 7.9, 2.6, 7.9, ""), (2.6, 7.9, 2.6, 10.4, ""),
        (2.6, 10.4, 3.7, 10.4, ""),  # loop back to diamond
        (6.3, 10.4, 6.3, 7.0, "нет/yo‘q"), (6.3, 7.0, 5, 7.0, ""),
        (5, 6.2, 5, 5.85, ""), (5, 5.15, 5, 4.7, ""),
        (5, 3.5, 5, 3.2, "да/ha"), (6.4, 4.1, 7.1, 4.1, "нет/yo‘q"),
        (8.4, 3.7, 8.4, 2.8, ""), (8.4, 2.8, 6.8, 2.8, ""),
        (5, 2.4, 5, 2.15, ""), (5, 1.45, 5, 1.15, ""), (5, 0.45, 5, 0.1, ""),
    ]
    return _gost_flowchart(N, A, (6.2, 11.5))


def _classic_fractal_image():
    """GOST flowchart of fractal construction (chaos game / IFS)."""
    N = [
        ("start", 5, 13.2, 2.6, 0.8, "Начало / Boshlash"),
        ("io", 5, 12.0, 4.6, 0.8, "Ввод: IFS {Tᵢ, pᵢ}"),
        ("proc", 5, 10.8, 3.6, 0.8, "P=(0.5,0.5); i=0"),
        ("cond", 5, 9.5, 2.4, 1.1, "i < N ?"),
        ("proc", 5, 8.1, 4.6, 0.8, "выбрать Tₖ по pₖ"),
        ("proc", 5, 7.0, 3.4, 0.7, "P = Aₖ·P + bₖ"),
        ("cond", 5, 5.7, 2.6, 1.1, "i > warmup ?"),
        ("proc", 8.0, 5.7, 2.8, 0.8, "нанести P"),
        ("proc", 5, 4.4, 2.4, 0.7, "i = i + 1"),
        ("proc", 5, 3.0, 4.4, 0.8, "растеризация → воксели"),
        ("proc", 5, 1.9, 3.8, 0.7, "marching cubes → 3D"),
        ("io", 5, 0.8, 3.8, 0.7, "Вывод: 3D-модель"),
        ("start", 5, -0.3, 2.6, 0.8, "Конец / Tugadi"),
    ]
    A = [
        (5, 12.8, 5, 12.4, ""), (5, 11.6, 5, 11.2, ""), (5, 10.4, 5, 10.05, ""),
        (5, 8.95, 5, 8.5, "да/ha"), (5, 7.7, 5, 7.35, ""), (5, 6.65, 5, 6.25, ""),
        (5, 5.15, 5, 4.75, "нет/yo‘q"),  # i>warmup? no → i=i+1
        (6.3, 5.7, 6.6, 5.7, "да/ha"), (8.0, 5.3, 8.0, 4.4, ""), (8.0, 4.4, 6.2, 4.4, ""),
        (3.8, 4.4, 2.6, 4.4, ""), (2.6, 4.4, 2.6, 9.5, ""), (2.6, 9.5, 3.8, 9.5, ""),  # loop back
        (6.2, 9.5, 6.2, 3.4, "нет/yo‘q"), (6.2, 3.4, 5, 3.4, ""),
        (5, 2.65, 5, 2.25, ""), (5, 1.55, 5, 1.15, ""), (5, 0.45, 5, 0.1, ""),
    ]
    return _gost_flowchart(N, A, (6.4, 10.0))


def build_algorithm_docx() -> bytes:
    """Build the bilingual DCNF-CRF algorithm Word document; return .docx bytes."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    title = doc.add_heading("Алгоритм DCNF CRF · DCNF CRF algoritmi", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        "Liu et al. «Deep Convolutional Neural Fields for Depth Estimation», "
        "CVPR 2015 + наши добавления (3 similarity, DA V2 unary, guided filter).")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in sub.runs:
        r.italic = True
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # --- Diagrams: architecture + two GOST flowcharts ---
    doc.add_heading("1. Блок-схемы · Blok-sxemalar", level=1)

    def _add_diagram(subheading, img):
        doc.add_heading(subheading, level=2)
        if img is not None:
            doc.add_picture(img, width=Inches(5.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_diagram("1.1 Архитектура (Liu et al.) · Arxitektura", _pipeline_image())
    _add_diagram("1.2 Блок-схема: Depth (CRF) · Depth (CRF) blok-sxema",
                 _classic_depth_image())
    _add_diagram("1.3 Блок-схема: построение фрактала · Fraktal qurish blok-sxema",
                 _classic_fractal_image())

    # --- Algorithm steps ---
    doc.add_heading("2. Пошаговый алгоритм · Bosqichma-bosqich algoritm", level=1)
    for i, (tru, tuz, formula, eru, euz) in enumerate(_STEPS, 1):
        p = doc.add_paragraph()
        rb = p.add_run(f"Шаг {i} · Qadam {i}: {tru} · {tuz}")
        rb.bold = True
        fp = doc.add_paragraph()
        _mono(fp.add_run(formula))
        doc.add_paragraph(f"RU: {eru}")
        doc.add_paragraph(f"UZ: {euz}")

    # --- Metrics table ---
    doc.add_heading("3. Метрики оценки · Baholash metrikalari", level=1)
    doc.add_paragraph(
        "Метрики из Liu et al. (Tables 1–3) требуют истинную глубину (ground truth). "
        "Tegishli metrikalar haqiqiy chuqurlikni (ground truth) talab qiladi.")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for c, txt in zip(hdr, ["Метрика", "Формула", "Пояснение (RU/UZ)", "Лучше"]):
        c.paragraphs[0].add_run(txt).bold = True
    for sym, nru, nuz, formula, eru, euz, better in _METRICS:
        cells = table.add_row().cells
        cells[0].text = f"{sym}\n{nru} · {nuz}"
        _mono(cells[1].paragraphs[0].add_run(formula))
        cells[2].text = f"RU: {eru}\nUZ: {euz}"
        cells[3].text = better

    doc.add_heading("4. Результаты · Natijalar", level=1)
    doc.add_paragraph(
        "В приложении ground-truth глубины нет, поэтому используются ОТНОСИТЕЛЬНЫЕ "
        "метрики между методами (edge alignment, useful detail, smoothness, depth "
        "range) для сравнения Unary / Make3D / Full CRF / Depth Anything V2. "
        "Ilovada ground-truth yo‘q, shuning uchun usullar orasidagi NISBIY metrikalar "
        "ishlatiladi.")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
